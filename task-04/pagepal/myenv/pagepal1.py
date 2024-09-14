from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import  Application,CommandHandler, filters, ContextTypes, MessageHandler, CallbackQueryHandler, ConversationHandler
from docx import Document
from docx.shared import Pt
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
import requests
import csv
from dotenv import load_dotenv
import os
import logging
import json
#defaults = Defaults(timeout=60)
# Define states for conversation
STEP1, STEP2, STEP3, STEP4 = range(4)

# Load environment variables
load_dotenv()
TOKEN = os.getenv("TOKEN")
C_ID = os.getenv("C_ID")
C_SEC = os.getenv("C_SEC")

logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO,
    handlers=[
        logging.FileHandler("bot.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Google Books API credentials setup
SCOPES = ['https://www.googleapis.com/auth/books']
creds = Credentials(
    None,
    refresh_token=TOKEN,
    client_id=C_ID,
    client_secret=C_SEC,
    token_uri='https://oauth2.googleapis.com/token',
    scopes=SCOPES
)

# Refresh the token if necessary
if creds.expired and creds.refresh_token:
    creds.refresh(Request())
access_token = creds.token


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Welcome To PagePal: Your New Reading Companion!")


async def book(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Which genre of books are you looking for?")
    print('hiiii')
    return STEP1


async def STEP1(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_msg = update.message.text
    response = requests.get(f'https://www.googleapis.com/books/v1/volumes?q={user_msg}')
  
    if response.status_code == 200:
        logger.info("SEARCH RESULTS RECEIVED")
        data = response.json()
        h = ['Book Title', 'Author', 'Description', 'OG Date of Publication', 'Language', 'Preview Link']
        try:
            with open("booksearch.csv", 'w', newline='') as f:
                fp = csv.writer(f, delimiter=',')
                fp.writerow(h)
                for item in data.get('items', []):
                    d = [item["volumeInfo"].get("title", "N/A"),
                         ", ".join(item["volumeInfo"].get("authors", [])),
                         item["volumeInfo"].get("description", "N/A"),
                         item["volumeInfo"].get("publishedDate", "N/A"),
                         item["volumeInfo"].get("language", "N/A"),
                         item["volumeInfo"].get("previewLink", "N/A")]
                    fp.writerow(d)

            chat_id = update.message.chat_id
            await context.bot.send_document(chat_id=chat_id, document=open("booksearch.csv", 'rb'))
            await update.message.reply_text("Happy Reading! :)")
        except Exception as e:
            logger.error(f"Error writing to CSV or sending the file: {e}")
            await update.message.reply_text("An error occurred while processing your request. Please try again later.")
    else:
        await update.message.reply_text("Sorry, no related results found.")
    
    return ConversationHandler.END


async def preview(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logger.info("Entered preview function")
    await update.message.reply_text("Type the name of the book for its preview link")
    return STEP2


async def STEP2(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_msg = update.message.text
    response = requests.get(f'https://www.googleapis.com/books/v1/volumes?q={user_msg}')
    print(f'user_msg')
    if response.status_code == 200:
        logger.info("preview link acquired")
        data =response.json()
        logger.info(data)
        print(f'user_msg')
        print(data)
        if data.get("items"):
            print('the book deatils recieved')
            first_book = data["items"][0]
            book_title = first_book["volumeInfo"].get("title", "Unknown Title")
            preview_link = first_book["volumeInfo"].get("previewLink", "No Preview Available")
            await update.message.reply_text(f"Here's the link to {book_title}: {preview_link}")
        else:
            await update.message.reply_text("No preview link found for the given title.")
    else:
        await update.message.reply_text("Failed to retrieve book details.")
    
    return ConversationHandler.END


async def help(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        """
        Hi there! Let me guide you through the utilities of this bot.
        
        /start : To start the conversation 
        /book  : To find the books from a certain genre
        /preview : To get the link to the preview of a book
        /list   : To make changes to / view your reading list
        /help   : For the list of commands I respond to"""
    ) 


async def list(update: Update, context: ContextTypes.DEFAULT_TYPE):
    print("Entered list function")
    await update.message.reply_text("Type in the name of a book:")
    return STEP3


async def STEP3(update: Update, context: ContextTypes.DEFAULT_TYPE):
    bookname = update.message.text
    context.user_data['bookname'] = bookname
    if context.user_data.get(bookname):
        print("booknames been saved")
        await update.message.reply_text("Now, type in 'reading list' to bring changes to your personal reading list.")
        return await readinglist(bookname, update, context)


async def readinglist(bookname, update: Update, context: ContextTypes.DEFAULT_TYPE):
    headers = {"Authorization": f"Bearer {access_token}"}
    shelflist = requests.get('https://www.googleapis.com/books/v1/mylibrary/bookshelves', headers=headers)
    
    if shelflist.status_code == 200:
        list_ = shelflist.json()
        if 'items' in list_:
            await update.message.reply_text("Which of these shelves do you want to bring changes to?")
            for item in list_['items']:
                await update.message.reply_text(f"* {item['title']} ?")
            
            shelf_id = list_['items'][0]['id']
            context.user_data['shelfid'] = shelf_id
            return await Response1(bookname, shelf_id, update, context) 
        else:
            await update.message.reply_text("You have no bookshelves available.")
            return ConversationHandler.END
    else:
        await update.message.reply_text("Failed to retrieve your bookshelves.")
        return ConversationHandler.END


async def Response1(bookname, shelfid, update: Update, context: ContextTypes.DEFAULT_TYPE):
    headers = {"Authorization": f"Bearer {access_token}"}
    shelflist = requests.get(f'https://www.googleapis.com/books/v1/mylibrary/bookshelves/{shelfid}', headers=headers)
    if shelflist.status_code == 200:
        list_ = shelflist.json()
        found_book = next((item for item in list_.get('items', []) if item["volumeInfo"]["title"].lower() == bookname.lower()), None)
        if found_book:
            vol_id = found_book['id']
            keyboard = [
                [InlineKeyboardButton('Add a Book', callback_data="addbook"),
                 InlineKeyboardButton('Delete A Book', callback_data='delbook'),
                 InlineKeyboardButton('View Reading List', callback_data='read_list')]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            await update.message.reply_text('Please choose an option:', reply_markup=reply_markup)

            context.user_data['shelfid'] = shelfid
            context.user_data['vol_id'] = vol_id
            return STEP4
        else:
            await update.message.reply_text("No volumes found in the selected shelf.")
            return ConversationHandler.END
    else:
        await update.message.reply_text("This shelf is empty.")
        return ConversationHandler.END


async def STEP4(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    data = query.data

    vol_id = context.user_data.get('vol_id')
    shelfid = context.user_data.get('shelfid')
    headers = {"Authorization": f"Bearer {access_token}"}
    params = {"volumeId": vol_id}
    
    if data == "addbook":
        action = requests.post(f'https://www.googleapis.com/books/v1/mylibrary/bookshelves/{shelfid}/addVolume', headers=headers, params=params)
        if action.status_code == 200:
            await query.edit_message_text("VOLUME ADDED SUCCESSFULLY!")
        else:
            await query.edit_message_text("Failed to add the volume.")
    
    elif data == 'delbook':
        action = requests.delete(f'https://www.googleapis.com/books/v1/mylibrary/bookshelves/{shelfid}/removeVolume', headers=headers, params=params)
        if action.status_code == 200:
            await query.edit_message_text("VOLUME DELETED SUCCESSFULLY!")
        else:
            await query.edit_message_text("Failed to delete the volume.")
    
    elif data == 'read_list':
        action = requests.get(f"https://www.googleapis.com/books/v1/mylibrary/bookshelves/{shelfid}/volumes", headers=headers)
        if action.status_code == 200:
            doc = Document()
            doc.add_heading('BOOKS FROM THE REQUESTED SHELF:', 0)
            obj = action.json()
            books = obj.get("items", [])
            for i, book in enumerate(books):
                para = doc.add_paragraph(f'{i + 1}. {book["volumeInfo"]["title"]}\n')
                para.style.font.size = Pt(12)
            doc.save("docc.docx")
            chat_id = update.effective_chat.id
            await context.bot.send_document(chat_id=chat_id, document=open("docc.docx", 'rb'))
        else:
            await query.edit_message_text("Failed to retrieve the reading list.")

    return ConversationHandler.END


def main() -> None:
    application = Application.builder().token(TOKEN).connect_timeout(5).build()

    convo_handler = ConversationHandler(
        entry_points=[CommandHandler('book', book)],
        states={
            STEP1: [MessageHandler(filters.TEXT & ~filters.COMMAND, STEP1)],
            STEP2: [MessageHandler(filters.TEXT & ~filters.COMMAND, STEP2)],
            STEP3: [MessageHandler(filters.TEXT & ~filters.COMMAND, STEP3)],
            STEP4: [CallbackQueryHandler(STEP4)],
        },
        fallbacks=[]
    )

    application.add_handler(convo_handler)
    application.add_handler(CommandHandler('start', start))
    application.add_handler(CommandHandler('help', help))
    application.add_handler(CommandHandler('preview', preview))
    application.add_handler(CommandHandler('list', list))
    application.run_polling()

if __name__ == '__main__':
    import asyncio
    asyncio.run(main())
